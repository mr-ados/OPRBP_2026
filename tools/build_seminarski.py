# -*- coding: utf-8 -*-
"""Build the UFC seminar DOCX and a simple ER diagram image."""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
IMG = DOCS / "img"
OUTPUT = DOCS / "UFC_seminarski_rad.docx"
ER_IMAGE = IMG / "er_ufc.png"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_FILL = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, title in enumerate(headers):
        set_cell_text(hdr[i], title, bold=True)
        set_cell_shading(hdr[i], LIGHT_FILL)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["CodeBlock"]
    for line in code.strip().splitlines():
        run = p.add_run(line.rstrip() + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def draw_er_diagram() -> None:
    IMG.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1700, 1050), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path(r"C:\Windows\Fonts\arial.ttf")
    bold_path = Path(r"C:\Windows\Fonts\arialbd.ttf")
    font = ImageFont.truetype(str(font_path), 24) if font_path.exists() else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 18) if font_path.exists() else ImageFont.load_default()
    bold = ImageFont.truetype(str(bold_path), 24) if bold_path.exists() else font

    boxes = {
        "Country": (70, 110, 250, 185),
        "Region": (315, 110, 495, 185),
        "City": (560, 110, 740, 185),
        "Event": (805, 110, 985, 185),
        "Fight": (1050, 110, 1250, 185),
        "WeightClass": (1350, 110, 1600, 185),
        "FightFormat": (1350, 245, 1600, 320),
        "VictoryMethod": (1350, 380, 1600, 455),
        "Referee": (1350, 515, 1600, 590),
        "Stance": (70, 450, 250, 525),
        "Fighter": (315, 450, 515, 525),
        "CareerStats": (315, 620, 535, 695),
        "Participant": (670, 450, 900, 525),
        "Performance": (670, 620, 930, 695),
        "StrikeBreakdown": (670, 790, 980, 865),
    }

    def mid_y(name: str) -> int:
        x1, y1, x2, y2 = boxes[name]
        return (y1 + y2) // 2

    def mid_x(name: str) -> int:
        x1, y1, x2, y2 = boxes[name]
        return (x1 + x2) // 2

    def right(name: str) -> tuple[int, int]:
        x1, y1, x2, y2 = boxes[name]
        return x2, (y1 + y2) // 2

    def left(name: str) -> tuple[int, int]:
        x1, y1, x2, y2 = boxes[name]
        return x1, (y1 + y2) // 2

    def top(name: str) -> tuple[int, int]:
        x1, y1, x2, y2 = boxes[name]
        return (x1 + x2) // 2, y1

    def bottom(name: str) -> tuple[int, int]:
        x1, y1, x2, y2 = boxes[name]
        return (x1 + x2) // 2, y2

    def label(x: int, y: int, text: str) -> None:
        draw.rectangle((x - 26, y - 13, x + 26, y + 13), fill="white")
        draw.text((x - 18, y - 11), text, fill=(70, 70, 70), font=small)

    def hline(a: str, b: str, text: str) -> None:
        ax, ay = right(a)
        bx, by = left(b)
        draw.line((ax, ay, bx, by), fill=(80, 80, 80), width=3)
        label((ax + bx) // 2, ay - 18, text)

    def vline(a: str, b: str, text: str) -> None:
        ax, ay = bottom(a)
        bx, by = top(b)
        draw.line((ax, ay, bx, by), fill=(80, 80, 80), width=3)
        label(ax + 34, (ay + by) // 2, text)

    for a, b, text in [
        ("Country", "Region", "1:N"),
        ("Region", "City", "1:N"),
        ("City", "Event", "1:N"),
        ("Event", "Fight", "1:N"),
        ("Fight", "WeightClass", "N:1"),
        ("Fight", "FightFormat", "N:1"),
        ("Fight", "VictoryMethod", "N:1"),
        ("Fight", "Referee", "N:1"),
        ("Stance", "Fighter", "1:N"),
        ("Fighter", "Participant", "1:N"),
    ]:
        hline(a, b, text)

    vline("Fighter", "CareerStats", "1:1")
    vline("Participant", "Performance", "1:1")
    vline("Performance", "StrikeBreakdown", "1:1")

    # Orthogonal relation from fight to participant keeps the central M:N resolution readable.
    fx, fy = bottom("Fight")
    px, py = top("Participant")
    joint_y = 350
    draw.line((fx, fy, fx, joint_y, px, joint_y, px, py), fill=(80, 80, 80), width=3)
    label((fx + px) // 2, joint_y - 18, "1:N")

    for name, (x1, y1, x2, y2) in boxes.items():
        draw.rounded_rectangle((x1, y1, x2, y2), radius=12, fill=(232, 238, 245), outline=(31, 77, 120), width=3)
        text_width = draw.textlength(name, font=bold)
        draw.text((x1 + (x2 - x1 - text_width) / 2, y1 + 16), name, fill=(11, 37, 69), font=bold)

    draw.text((70, 35), "UFC_OPRBP - ER model (skraćeni prikaz)", fill=(11, 37, 69), font=bold)
    image.save(ER_IMAGE)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "CodeBlock" not in [s.name for s in styles]:
        code_style = styles.add_style("CodeBlock", 1)
    else:
        code_style = styles["CodeBlock"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Inches(0.25)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(6)

    return doc


def build_doc() -> None:
    draw_er_diagram()
    doc = setup_document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SVEUČILIŠTE U SPLITU\nPRIRODOSLOVNO-MATEMATIČKI FAKULTET")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("\n\n\n")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SEMINARSKI RAD\nODABRANA POGLAVLJA RELACIJSKIH BAZA PODATAKA\n\nUFC relacijska baza podataka")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_BLUE

    doc.add_paragraph("\n\n")
    meta = doc.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(meta.cell(0, 0), "Profesorica:", True)
    set_cell_text(meta.cell(0, 1), "Monika Mladenović")
    set_cell_text(meta.cell(1, 0), "Student:", True)
    set_cell_text(meta.cell(1, 1), "Mario")
    for row in meta.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph("\n\nSplit, svibanj 2026.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Sadržaj", level=1)
    for item in [
        "1 UVOD",
        "2 IZRADA MODELA",
        "3 RELACIJE",
        "4 LOGIČKI I RELACIJSKI MODEL",
        "5 IZRADA BAZE PODATAKA PREMA MODELU",
        "6 UPITI NA BAZU PODATAKA",
        "7 OBJEKTI S PREDAVANJA",
        "8 KORIŠTENO I NEKORIŠTENO S PREDAVANJA",
        "9 ZAKLJUČAK",
    ]:
        doc.add_paragraph(item)
    doc.add_page_break()

    doc.add_heading("1 UVOD", level=1)
    doc.add_heading("1.1 Općenito", level=2)
    doc.add_paragraph(
        "Ultimate Fighting Championship (UFC) je najpoznatija svjetska organizacija "
        "mješovitih borilačkih vještina. Organizacija se sastoji od evenata na kojima "
        "se održavaju borbe po težinama, a svaka borba ima dva natjecatelja, rezultat, "
        "metodu završetka i velik broj statističkih pokazatelja."
    )
    doc.add_heading("1.2 Opis projekta", level=2)
    doc.add_paragraph(
        "Cilj projekta je iz denormaliziranog Kaggle dataseta izgraditi relacijsku bazu "
        "podataka u SQL Serveru. Baza omogućuje evidenciju evenata, lokacija, boraca, "
        "borbi, pobjednika, sudaca, težinskih kategorija i statistike borbe. Podaci se "
        "prvo učitavaju u staging tablice, a zatim se transformiraju u normalizirani model."
    )

    doc.add_heading("2 IZRADA MODELA", level=1)
    doc.add_paragraph(
        "Model je namjerno podijeljen u više shema: `stg` za sirove CSV podatke, `geo` "
        "za lokacije, `ref` za šifarnike, `ufc` za domenske tablice i `audit` za zapis "
        "promjena nastalih triggerom."
    )
    add_table(
        doc,
        ["Entitet", "Opis"],
        [
            ["Country, Region, City", "Normalizirana lokacija UFC eventa."],
            ["Event", "UFC priredba s nazivom, datumom i lokacijom."],
            ["Fighter", "Borac s osnovnim biografskim podacima."],
            ["FighterCareerStats", "Karijerna statistika borca; 1:1 s Fighter."],
            ["Stance", "Stav borca: Orthodox, Southpaw, Switch i sl."],
            ["WeightClass", "Težinska kategorija, uključujući ženske, interim i catch weight kategorije."],
            ["FightFormat", "Broj predviđenih rundi."],
            ["VictoryMethod, VictoryDetail", "Način pobjede i dodatni opis, npr. Decision - Unanimous."],
            ["Referee", "Sudac borbe."],
            ["Fight", "Centralna tablica borbe."],
            ["FightParticipant", "Razrješava M:N vezu između borbe i borca."],
            ["FightPerformanceStats", "Ukupna statistika borca u pojedinoj borbi."],
            ["FightStrikeBreakdown", "Raspodjela udaraca po meti i poziciji."],
            ["ChangeLog", "Audit tablica koju puni trigger."],
        ],
    )

    doc.add_heading("3 RELACIJE", level=1)
    doc.add_paragraph(
        "U modelu su prikazane sve tri osnovne vrste relacija. Relacija 1:1 postoji između "
        "`Fighter` i `FighterCareerStats`, jer jedan borac ima jedan redak karijerne statistike. "
        "Relacija 1:N pojavljuje se između `Event` i `Fight`, jer jedan event ima više borbi. "
        "Relacija M:N između `Fighter` i `Fight` razrješava se tablicom `FightParticipant`, "
        "jer jedan borac može nastupiti u više borbi, a jedna borba ima više boraca."
    )

    doc.add_heading("4 LOGIČKI I RELACIJSKI MODEL", level=1)
    doc.add_paragraph("Na slici je prikazan skraćeni ER model s glavnim tablicama i relacijama.")
    doc.add_picture(str(ER_IMAGE), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Slika 1. Skraćeni ER model baze UFC_OPRBP").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(
        doc,
        ["Tablica", "Primarni kljuc", "Vazni strani kljucevi"],
        [
            ["ufc.Event", "event_id", "city_id -> geo.City"],
            ["ufc.Fight", "fight_id", "event_id, weight_class_id, victory_method_id, referee_id, winner_fighter_id"],
            ["ufc.Fighter", "fighter_id", "stance_id -> ref.Stance"],
            ["ufc.FightParticipant", "fight_id + corner_color", "fight_id -> ufc.Fight, fighter_id -> ufc.Fighter"],
            ["ufc.FightPerformanceStats", "fight_id + corner_color", "FK prema FightParticipant"],
            ["ufc.FightStrikeBreakdown", "fight_id + corner_color", "FK prema FightParticipant"],
        ],
    )

    doc.add_heading("5 IZRADA BAZE PODATAKA PREMA MODELU", level=1)
    doc.add_paragraph(
        "Izrada baze podataka organizirana je kroz SQL skripte. Prve skripte stvaraju bazu, "
        "sheme i staging tablice. Nakon učitavanja CSV datoteka, transformacijska skripta "
        "popunjava normalizirane tablice i koristi `TRY_CONVERT` kako bi se tekstualni CSV "
        "podaci pretvorili u datume, brojeve i logičke vrijednosti."
    )
    add_table(
        doc,
        ["Skripta", "Namjena"],
        [
            ["00_create_database.sql", "Stvara bazu UFC_OPRBP."],
            ["01_create_schema.sql", "Stvara sheme, tablice, ključeve i ograničenja."],
            ["02_create_staging.sql", "Stvara staging tablice prema CSV headerima."],
            ["03_load_or_import_notes.sql", "Učitava CSV datoteke pomoću BULK INSERT."],
            ["04_transform_to_model.sql", "Transformira staging podatke u relacijski model."],
            ["05_views_triggers_transactions.sql", "Stvara poglede i audit trigger."],
            ["06_procedures_json.sql", "Stvara procedure i JSON primjere."],
            ["07_demo_queries.sql", "Sadrzi upite za demonstraciju i provjeru."],
        ],
    )

    doc.add_heading("6 UPITI NA BAZU PODATAKA", level=1)
    doc.add_heading("6.1 Osnovni SELECT i JOIN", level=2)
    add_code(
        doc,
        """
        SELECT TOP (30)
            event_date, event_name, division_name,
            red_fighter, blue_fighter, winner, method
        FROM ufc.v_event_results
        WHERE country = N'USA'
        ORDER BY event_date DESC;
        """,
    )
    doc.add_heading("6.2 Grupiranje i HAVING", level=2)
    add_code(
        doc,
        """
        SELECT division_name, COUNT(*) AS fight_count
        FROM ufc.v_event_results
        GROUP BY division_name
        HAVING COUNT(*) >= 100
        ORDER BY fight_count DESC;
        """,
    )
    doc.add_heading("6.3 CTE i window funkcija", level=2)
    add_code(
        doc,
        """
        WITH wins_by_division AS
        (
            SELECT wc.division_name, fi.fighter_name, COUNT(*) AS wins
            FROM ufc.FightParticipant fp
            JOIN ufc.Fighter fi ON fi.fighter_id = fp.fighter_id
            JOIN ufc.Fight f ON f.fight_id = fp.fight_id
            JOIN ref.WeightClass wc ON wc.weight_class_id = f.weight_class_id
            WHERE fp.is_winner = 1
            GROUP BY wc.division_name, fi.fighter_name
        )
        SELECT *, DENSE_RANK() OVER
            (PARTITION BY division_name ORDER BY wins DESC) AS division_rank
        FROM wins_by_division;
        """,
    )

    doc.add_heading("7 OBJEKTI S PREDAVANJA", level=1)
    doc.add_paragraph("Projekt demonstrira gradivo s vjezbi kroz konkretne SQL Server objekte.")
    add_bullets(
        doc,
        [
            "Pogledi: `ufc.v_event_results`, `ufc.v_fighter_summary`, `ufc.v_weight_class_statistics`.",
            "Trigger: `ufc.trg_Fight_Audit` upisuje stare i nove vrijednosti u JSON obliku.",
            "Transakcije: `ufc.sp_update_fight_result` radi `COMMIT` samo ako je parametar `@commit_changes = 1`, inace radi `ROLLBACK`.",
            "Procedure: dohvat borbi, usporedba boraca, paging rezultata i promjena rezultata.",
            "JSON: `FOR JSON PATH`, `ISJSON` i `JSON_VALUE` koriste se za izvoz i uvoz rezultata borbe.",
        ],
    )
    add_code(
        doc,
        """
        EXEC ufc.sp_fights_paging
            @division_name = N'lightweight',
            @skip = 0,
            @getRows = 10;

        EXEC ufc.sp_get_event_card_json
            @event_id = '6e380a4d73ab4f0e';
        """,
    )

    doc.add_heading("8 KORIŠTENO I NEKORIŠTENO S PREDAVANJA", level=1)
    add_table(
        doc,
        ["Tema", "Status", "Obrazlozenje"],
        [
            ["SQL DDL/DML/DQL", "Korišteno", "Baza se stvara skriptama, podaci se uvoze, ažuriraju i dohvaćaju SQL upitima."],
            ["JOIN i agregacije", "Korišteno", "Koriste se u pogledima i demo upitima za statistiku borbi."],
            ["Pogledi", "Korišteno", "Složeni upiti zapakirani su u viewove."],
            ["Triggeri", "Korišteno", "Audit trigger zapisuje promjene nad tablicom Fight."],
            ["Transakcije", "Korišteno", "Procedure imaju rollback/commit scenarij."],
            ["Procedure", "Korišteno", "Parametrizirani dohvat i promjena rezultata."],
            ["JSON", "Korišteno", "Izvoz event carda i uvoz promjene rezultata."],
            ["APEX", "Nije korišteno", "Nije tražen i vezan je uz Oracle, a projekt je u SQL Serveru."],
            ["MongoDB/NoSQL", "Nije korišteno", "Tema projekta je relacijska baza podataka."],
            ["Distribuirane baze", "Nije korišteno", "Lokalni studentski projekt nema potrebu za distribuiranom arhitekturom."],
        ],
    )

    doc.add_heading("9 ZAKLJUČAK", level=1)
    doc.add_paragraph(
        "Projekt pokazuje kako se veliki denormalizirani sportski dataset može pretvoriti "
        "u relacijsku bazu podataka. Staging sloj čuva izvorni oblik podataka, dok relacijski "
        "model omogućuje jasne veze, primarne i strane ključeve, ograničenja i pouzdanije "
        "upite. Kroz poglede, triggere, transakcije, procedure i JSON demonstrirane su teme "
        "obrađene na vježbama."
    )

    doc.add_page_break()
    doc.add_heading("Dodatak: testni scenariji", level=1)
    add_bullets(
        doc,
        [
            "Broj redaka u staging tablicama uspoređuje se s brojem redaka u CSV datotekama.",
            "Svaka borba treba imati dva zapisa u `ufc.FightParticipant` kad dataset ima oba borca.",
            "Strani ključevi sprječavaju orphan zapise.",
            "CHECK ograničenja sprječavaju nelogične statistike poput landed > attempted.",
            "JSON procedure se provjeravaju funkcijom `ISJSON`.",
            "Rollback demo potvrđuje da se transakcija može poništiti bez trajnih promjena.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "UFC_OPRBP - seminarski rad"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {ER_IMAGE}")


if __name__ == "__main__":
    build_doc()
